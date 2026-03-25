"""
기술서 파일 텍스트 추출
지원 형식: .docx, .pdf, .txt, .xlsx, .xls, .csv
"""
from pathlib import Path
from utils.logger import get_logger

log = get_logger('doc_parser')


def extract_text(file_bytes: bytes, filename: str) -> str:
    """업로드된 파일에서 텍스트를 추출한다."""
    ext = Path(filename).suffix.lower()

    try:
        if ext == '.txt':
            return _parse_txt(file_bytes)
        elif ext == '.docx':
            return _parse_docx(file_bytes)
        elif ext == '.pdf':
            return _parse_pdf(file_bytes)
        elif ext in ('.xlsx', '.xls'):
            return _parse_excel(file_bytes)
        elif ext == '.csv':
            return _parse_csv(file_bytes)
        else:
            log.warning('지원하지 않는 형식: %s — 텍스트로 시도', ext)
            return _parse_txt(file_bytes)
    except Exception as e:
        log.error('파일 파싱 실패 [%s]: %s', filename, e)
        return ''


def _parse_txt(file_bytes: bytes) -> str:
    for enc in ['utf-8', 'cp949', 'euc-kr', 'latin-1']:
        try:
            return file_bytes.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return file_bytes.decode('utf-8', errors='replace')


def _parse_docx(file_bytes: bytes) -> str:
    from io import BytesIO
    from docx import Document
    doc = Document(BytesIO(file_bytes))
    lines = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            lines.append(text)
    # 테이블 내용도 추출
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(' | '.join(cells))
    return '\n'.join(lines)


def _parse_excel(file_bytes: bytes) -> str:
    """Excel(.xlsx/.xls) → 시트별 테이블 텍스트"""
    from io import BytesIO
    import pandas as pd

    xls = pd.ExcelFile(BytesIO(file_bytes))
    lines = []
    for sheet_name in xls.sheet_names:
        df = xls.parse(sheet_name).fillna('')
        if df.empty:
            continue
        lines.append(f'[시트: {sheet_name}]')
        # 헤더
        lines.append(' | '.join(str(c) for c in df.columns))
        # 데이터 (최대 200행)
        for _, row in df.head(200).iterrows():
            cells = [str(v).strip() for v in row if str(v).strip()]
            if cells:
                lines.append(' | '.join(cells))
        lines.append('')
    return '\n'.join(lines)


def _parse_csv(file_bytes: bytes) -> str:
    """CSV → 테이블 텍스트"""
    from io import BytesIO, StringIO
    import pandas as pd

    # 인코딩 자동 감지
    for enc in ['utf-8', 'cp949', 'euc-kr']:
        try:
            text = file_bytes.decode(enc)
            df = pd.read_csv(StringIO(text)).fillna('')
            break
        except (UnicodeDecodeError, Exception):
            continue
    else:
        return file_bytes.decode('utf-8', errors='replace')

    lines = [' | '.join(str(c) for c in df.columns)]
    for _, row in df.head(200).iterrows():
        cells = [str(v).strip() for v in row if str(v).strip()]
        if cells:
            lines.append(' | '.join(cells))
    return '\n'.join(lines)


def _parse_pdf(file_bytes: bytes) -> str:
    """PDF 텍스트 추출 — PyPDF2 또는 pdfplumber 사용"""
    from io import BytesIO
    try:
        import pdfplumber
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            pages = []
            for page in pdf.pages[:30]:  # 최대 30페이지
                text = page.extract_text()
                if text:
                    pages.append(text)
            return '\n'.join(pages)
    except ImportError:
        pass

    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(BytesIO(file_bytes))
        pages = []
        for page in reader.pages[:30]:
            text = page.extract_text()
            if text:
                pages.append(text)
        return '\n'.join(pages)
    except ImportError:
        log.error('PDF 파싱 라이브러리 없음 — pip install pdfplumber 또는 PyPDF2')
        return ''
